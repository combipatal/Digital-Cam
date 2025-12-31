import sys
import subprocess
import importlib

print("Starting script...")

def install(package):
    try:
        importlib.import_module(package)
        print(f"{package} is already installed.")
    except ImportError:
        print(f"Installing {package}...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package, "--no-input"])
            print(f"Successfully installed {package}.")
        except subprocess.CalledProcessError as e:
            print(f"Failed to install {package}: {e}")
            sys.exit(1)

# Ensure dependencies are installed BEFORE importing them
install('markdown')
install('python-docx') # package name is python-docx, import name is docx

try:
    import markdown
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    print("Imports successful.")
except ImportError as e:
    print(f"Import failed: {e}")
    sys.exit(1)

def md_to_docx(md_file, docx_file):
    print(f"Reading {md_file}...")
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except FileNotFoundError:
        print(f"File not found: {md_file}")
        return

    doc = Document()
    
    # Title
    title = doc.add_heading('FPGA 기반 실시간 이미지 프로세싱 시스템', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)

    try:
        doc.save(docx_file)
        print(f"Successfully created {docx_file}")
    except PermissionError:
        print(f"Permission denied: {docx_file}. Is the file open?")

if __name__ == "__main__":
    md_path = r"c:\git\Digital_Cam\PORTFOLIO.md"
    docx_path = r"c:\git\Digital_Cam\PORTFOLIO.docx"
    
    try:
        md_to_docx(md_path, docx_path)
    except Exception as e:
        print(f"Error: {e}")
